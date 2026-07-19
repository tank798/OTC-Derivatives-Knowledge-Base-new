from __future__ import annotations

from pathlib import Path
import re
import tempfile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn

from models import ParsedDocument, SourceBlock
from utils.metadata import infer_metadata
from utils.text import clean_text, compact, is_page_number, is_toc_field, markdown_table, strip_repeated_front_structure


FRONT_PAGE_TIMESTAMP_RE = re.compile(r"^时间\s*[:：]\s*\d{4}[-/.]\d{1,2}[-/.]\d{1,2}$")


def open_word_document(path: Path):
    """Open OOXML Word files without changing the original container.

    Some exchange downloads use a .docx name while declaring the main part as
    macro-enabled Word content. python-docx rejects that content type even
    though the visible WordprocessingML is otherwise standard. For extraction,
    create a temporary macro-free package declaration and leave the source
    bytes untouched.
    """
    try:
        return Document(path), None
    except ValueError as exc:
        if "macroEnabled.main+xml" not in str(exc):
            raise
    temp_dir = tempfile.TemporaryDirectory(prefix="regulatory_docm_")
    converted = Path(temp_dir.name) / f"{path.stem}.docx"
    macro_type = b"application/vnd.ms-word.document.macroEnabled.main+xml"
    standard_type = b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"
    with ZipFile(path) as source, ZipFile(converted, "w", ZIP_DEFLATED) as target:
        for item in source.infolist():
            payload = source.read(item.filename)
            if item.filename == "[Content_Types].xml":
                payload = payload.replace(macro_type, standard_type)
            target.writestr(item, payload)
    return Document(converted), temp_dir


def strip_front_page_metadata(blocks: list[SourceBlock]) -> tuple[list[SourceBlock], int]:
    """Remove exact website-style timestamps from the Word front matter.

    Dates inside articles remain untouched.  The narrow label + ISO-like date
    pattern and first-ten-block scope distinguish page metadata from legal text.
    """
    removed = 0
    result: list[SourceBlock] = []
    for index, block in enumerate(blocks):
        if index < 10 and FRONT_PAGE_TIMESTAMP_RE.fullmatch(clean_text(block.text)):
            removed += 1
            continue
        result.append(block)
    return result, removed


def paragraph_xml_text(paragraph: Paragraph) -> str:
    """Read visible text from all WordprocessingML containers in a paragraph.

    ``Paragraph.text`` only walks direct runs and silently drops text nested in
    ``w:smartTag``, content controls and some field containers.  Regulatory
    documents use those containers for dates, tenors and numbered items, so the
    omission can change legal or financial meaning.
    """
    pieces: list[str] = []
    for element in paragraph._p.iter():
        if element.tag == qn("w:t"):
            if element.text:
                pieces.append(element.text)
        elif element.tag == qn("w:tab"):
            pieces.append("\t")
        elif element.tag in {qn("w:br"), qn("w:cr")}:
            pieces.append("\n")
    return clean_text("".join(pieces))


def table_cell_text(cell) -> str:
    values: list[str] = []
    for child in cell.iter_inner_content():
        if isinstance(child, Paragraph):
            value = paragraph_xml_text(child)
            if value:
                values.append(value)
        elif isinstance(child, Table):
            for row in child.rows:
                values.append(" | ".join(table_cell_text(nested) for nested in row.cells))
    return "\n".join(values)


def word_table_data(table: Table, table_index: int) -> tuple[list[list[str]], dict]:
    """Preserve Word grid spans and vertical merges without proxy duplication."""

    width = len(table.columns)
    matrix: list[list[str]] = []
    cells: list[dict] = []
    vertical_starts: dict[int, dict] = {}
    explicit_header_rows = 0
    warnings: list[str] = []
    for row_index, row in enumerate(table.rows):
        values = [""] * width
        proxies = list(row.cells)
        column = 0
        seen_in_row: set[int] = set()
        while column < min(width, len(proxies)):
            cell = proxies[column]
            cell_key = id(cell._tc)
            run = 1
            while column + run < len(proxies) and id(proxies[column + run]._tc) == cell_key:
                run += 1
            tc_pr = cell._tc.tcPr
            grid_span = getattr(getattr(tc_pr, "gridSpan", None), "val", None)
            colspan = max(1, int(grid_span or run))
            v_merge = getattr(tc_pr, "vMerge", None)
            v_value = getattr(v_merge, "val", None) if v_merge is not None else None
            continuation = v_merge is not None and v_value not in {"restart"}
            if continuation and column in vertical_starts:
                vertical_starts[column]["rowspan"] += 1
            elif cell_key not in seen_in_row:
                text = table_cell_text(cell)
                values[column] = text
                cell_row = {
                    "row": row_index,
                    "column": column,
                    "text": text,
                    "rowspan": 1,
                    "colspan": min(colspan, width - column),
                }
                cells.append(cell_row)
                if v_merge is not None:
                    vertical_starts[column] = cell_row
            seen_in_row.add(cell_key)
            column += max(run, colspan)
        matrix.append(values)
        tr_pr = row._tr.trPr
        if tr_pr is not None and tr_pr.find(qn("w:tblHeader")) is not None:
            explicit_header_rows += 1

    header_count = explicit_header_rows or (1 if matrix else 0)
    if not explicit_header_rows:
        warnings.append("Word未标记重复表头，按首行作为检索表头")
    data = {
        "table_id": f"table_word_{table_index}",
        "caption": "",
        "headers": matrix[:header_count],
        "rows": matrix[header_count:],
        "cells": cells,
        "column_count": width,
        "source_page_start": 0,
        "source_page_end": 0,
        "source_fragments": [{"table_index": table_index}],
        "parsing_warnings": warnings,
    }
    return matrix, data


def strip_word_toc(blocks: list[SourceBlock]) -> tuple[list[SourceBlock], int]:
    marker = next((index for index, block in enumerate(blocks[:20]) if re.sub(r"\s+", "", clean_text(block.text)) in {"目录", "目次"}), None)
    if marker is None:
        return blocks, 0
    # Word生成的目录通常有明确的TOC 1/TOC 2样式。优先按样式删除连续
    # 目录条目，不要依赖正文必须以“第一条”开头；业务指南、操作手册等
    # 经常采用“章 + 一、二、三”的结构，首个“第X条”可能直到附件合同
    # 才出现。旧逻辑会因此把整篇主指南误当目录删除。
    styled_end = marker + 1
    while styled_end < len(blocks) and re.match(r"^toc(?:\s*\d+)?$", clean_text(blocks[styled_end].style), re.I):
        styled_end += 1
    if styled_end > marker + 1:
        return blocks[:marker] + blocks[styled_end:], styled_end - marker
    # 业务指南常见无TOC样式的纯文本目录，条目末尾直接跟页码，
    # 正文从再次出现的“说明及声明”或第一章开始。按规范化标题寻找
    # 重复起点，不要求正文一定含“第X条”。
    def toc_key(value: str) -> str:
        return re.sub(r"\s+\d{1,4}\s*$", "", re.sub(r"\s+", " ", clean_text(value))).strip()

    guide_prefix = blocks[marker + 1:min(marker + 130, len(blocks))]
    page_suffixed = sum(
        bool(re.search(r"\s+\d{1,4}\s*$", clean_text(block.text)))
        and bool(re.match(r"^(?:第.+[章节]|[一二三四五六七八九十百]+[、.]|附件|附录|说明及声明)", clean_text(block.text)))
        for block in guide_prefix
    )
    if guide_prefix and page_suffixed >= 3:
        first_key = toc_key(guide_prefix[0].text)
        duplicate = next(
            (
                index for index in range(marker + 2, len(blocks))
                if toc_key(blocks[index].text) == first_key
                and not re.search(r"\s+\d{1,4}\s*$", clean_text(blocks[index].text))
            ),
            None,
        )
        if duplicate is not None:
            return blocks[:marker] + blocks[duplicate:], duplicate - marker
    for index in range(marker + 2, len(blocks)):
        value = re.sub(r"\s+", "", clean_text(blocks[index].text))
        earlier = {re.sub(r"\s+", "", clean_text(block.text)) for block in blocks[marker + 1:index]}
        following = [clean_text(block.text) for block in blocks[index + 1:index + 6]]
        if value in earlier and any(re.match(r"^第\s*[一二三四五六七八九十百千万零〇两\d ]+\s*条", item) for item in following):
            return blocks[:marker] + blocks[index:], index - marker
    # 部分Word目录没有重复章节标题，只有目录域生成的条目。此时从第一条
    # 正文向前保留连续的编/章/节标题，目录本身及页码条目全部丢弃。
    article_index = next(
        (index for index in range(marker + 1, len(blocks)) if re.match(r"^第\s*[一二三四五六七八九十百千万零〇两\d ]+\s*条", clean_text(blocks[index].text))),
        None,
    )
    if article_index is not None:
        start = article_index
        for index in range(article_index - 1, marker, -1):
            value = clean_text(blocks[index].text)
            if re.match(r"^第\s*[一二三四五六七八九十百千万零〇两\d ]+\s*[编篇部分章节](?:\s+.{0,80})?$", value) and not re.search(r"\.{2,}|[…·]{2,}|\s\d{1,4}$", value):
                start = index
                continue
            break
        return blocks[:marker] + blocks[start:], start - marker
    # 若Word目录域本身没有导出任何条目，而后面紧接正式标题和
    # “说明及声明”，仅删除孤立的“目录”标记。
    if any(compact(block.text) in {"说明及声明", "声明"} for block in blocks[marker + 1:marker + 6]):
        return blocks[:marker] + blocks[marker + 1:], 1
    return blocks, 0


def parse_docx(path: Path) -> ParsedDocument:
    document, temporary_package = open_word_document(path)
    blocks: list[SourceBlock] = []
    sequence = 0
    table_sequence = 0

    def append_item(item) -> None:
        nonlocal sequence, table_sequence
        sequence += 1
        if isinstance(item, Paragraph):
            text = paragraph_xml_text(item)
            if not text:
                return
            style = item.style.name if item.style else ""
            blocks.append(SourceBlock(text=text, style=style, source_kind="paragraph", block_id=f"b{sequence:05d}"))
        elif isinstance(item, Table):
            # WPS/网页导出的法规有时把整篇正文包在一列嵌套表格里。
            # 逐单元格递归可保留条款段落；真正的多列表格仍转Markdown。
            if len(item.columns) == 1:
                seen_cells: set[int] = set()
                for row in item.rows:
                    cell = row.cells[0]
                    cell_key = id(cell._tc)
                    if cell_key in seen_cells:
                        continue
                    seen_cells.add(cell_key)
                    for child in cell.iter_inner_content():
                        append_item(child)
                return
            table_sequence += 1
            rows, table_data = word_table_data(item, table_sequence)
            text = markdown_table(rows)
            if text:
                blocks.append(SourceBlock(
                    text=text,
                    style="Table",
                    source_kind="table",
                    block_id=f"b{sequence:05d}",
                    table_data=table_data,
                    parsing_warnings=list(table_data["parsing_warnings"]),
                ))

    for item in document.iter_inner_content():
        append_item(item)
    warnings: list[str] = []
    blocks, removed_toc = strip_word_toc(blocks)
    if removed_toc:
        warnings.append(f"已过滤{removed_toc}个Word目录段落")
    blocks, removed_front_metadata = strip_front_page_metadata(blocks)
    if removed_front_metadata:
        warnings.append(f"已过滤{removed_front_metadata}个Word首页时间元数据段落")
    before_noise = len(blocks)
    blocks = [block for block in blocks if not is_page_number(block.text) and not is_toc_field(block.text)]
    if len(blocks) != before_noise:
        warnings.append(f"已过滤{before_noise - len(blocks)}个Word页码或目录域段落")
    blocks, removed_front_structure = strip_repeated_front_structure(blocks)
    if removed_front_structure:
        warnings.append(f"已过滤{removed_front_structure}个Word前置目录标题")
    header_footer_text = []
    for section in document.sections:
        for container in (section.header, section.footer):
            header_footer_text.extend(clean_text(p.text) for p in container.paragraphs if clean_text(p.text))
    if header_footer_text:
        warnings.append("已识别Word页眉页脚，未混入正文：" + " | ".join(dict.fromkeys(header_footer_text))[:300])
    if temporary_package is not None:
        warnings.append("宏启用Word容器在临时目录转换声明后只读解析，原文件未修改")
    metadata = infer_metadata(blocks, path)
    if temporary_package is not None:
        temporary_package.cleanup()
    return ParsedDocument(path, "docx", blocks, metadata, warnings)

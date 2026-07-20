from pathlib import Path
import sys
import tempfile
import unittest

PROGRAM_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROGRAM_ROOT))

from models import ParsedDocument, SourceBlock
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from parsers.docx_parser import paragraph_xml_text, strip_front_page_metadata, strip_word_toc
from parsers.dispatcher import parse_file
from parsers.legacy_doc_parser import _blocks_from_plain_text
from parsers.text_parser import _OfficialBodyHTML
from parsers.pdf_parser import join_pdf_lines, merge_cross_page_paragraphs, normalize_pdf_table_cell, normalize_semantic_table, remove_toc_entries, split_semantic_table_content
from parsers.pdf_formula_overrides import apply_verified_formula_overrides
from parsers.pdf_content_overrides import apply_verified_pdf_content_overrides
from parsers.legacy_doc_formula_overrides import apply_verified_formula_overrides as apply_verified_doc_formula_overrides
from utils.metadata import infer_metadata
from utils.front_matter import clean_front_matter
from utils.structured import structured_blocks
from utils.text import clean_text, is_page_number


class MetadataAndPdfTests(unittest.TestCase):
    def test_pdf_wrapped_numbered_items_are_joined_without_promoting_the_lead_to_heading(self):
        lines = [
            "（一）向《私募办法》规定的合格投资者之外的单位、个人",
            "募集资金或者为投资者提供多人拼凑、资金借贷等满足合格投资者要求的便利；",
            "（二）通过报刊、电台、电视、互联网等公众传播媒体，讲",
            "座、报告会、分析会等方式向不特定对象宣传推介；",
        ]
        joined = join_pdf_lines(lines)
        self.assertEqual(
            joined,
            [
                "(一)向《私募办法》规定的合格投资者之外的单位、个人募集资金或者为投资者提供多人拼凑、资金借贷等满足合格投资者要求的便利;",
                "(二)通过报刊、电台、电视、互联网等公众传播媒体,讲座、报告会、分析会等方式向不特定对象宣传推介;",
            ],
        )

    def test_pdf_wrapped_item_can_continue_before_an_article_like_phrase(self):
        joined = join_pdf_lines([
            "（三）不符合本规定第六条第一款第（一）项至第（八）项、",
            "第六条第一款第（十）项、第七条的，可以依法处理；",
            "第七条 这是下一条正文。",
        ])
        self.assertEqual(
            joined,
            [
                "(三)不符合本规定第六条第一款第(一)项至第(八)项、第六条第一款第(十)项、第七条的,可以依法处理;",
                "第七条 这是下一条正文。",
            ],
        )

    def test_verified_attachment_cleanup_is_bounded_and_keeps_following_body(self):
        blocks = [
            SourceBlock("正文第一条。", block_id="b1"),
            SourceBlock("附件:1.场外衍生品报告内容与格式模板", block_id="b2"),
            SourceBlock("2.非公开发行公司债券备案内容与格式模板", block_id="b3"),
            SourceBlock("3.收益凭证报告内容与格式模板", block_id="b4"),
            SourceBlock("4.场外证券销售业务报告内容与格式模板", block_id="b5"),
            SourceBlock("5.证券公司登记、托管和结算报告内容与格式模板", block_id="b6"),
            SourceBlock("中国证券业协会", block_id="b7"),
            SourceBlock("2017年5月22日", block_id="b8"),
            SourceBlock("正文附件后的条款。", block_id="b9"),
        ]
        kept, descriptions = apply_verified_pdf_content_overrides(
            Path("关于加强场外衍生品业务自律管理的通知.docx"), blocks,
        )
        self.assertEqual([block.block_id for block in kept], ["b1", "b9"])
        self.assertEqual(len(descriptions), 1)

    def test_verified_attachment_cleanup_does_not_apply_to_unknown_file(self):
        blocks = [SourceBlock("附件:1.模板", block_id="b1"), SourceBlock("正文", block_id="b2")]
        kept, descriptions = apply_verified_pdf_content_overrides(Path("其他通知.pdf"), blocks)
        self.assertEqual([block.block_id for block in kept], ["b1", "b2"])
        self.assertEqual(descriptions, [])

    def test_structural_front_matter_cleaning_is_scoped_before_articles(self):
        document = ParsedDocument(
            Path("公开募集证券投资基金投资信用衍生品指引.docx"),
            "docx",
            [
                SourceBlock("公开募集证券投资基金投资信用衍生品指引", block_id="b1"),
                SourceBlock("证监会公告〔2019〕第1号", block_id="b2"),
                SourceBlock("现公布《公开募集证券投资基金投资信用衍生品指引》，自公布之日起施行。", block_id="b3"),
                SourceBlock("公开募集证券投资基金投资信用衍生品指引", block_id="b4"),
                SourceBlock("第一条 为规范基金投资行为，制定本指引。", block_id="b5"),
                SourceBlock("第十一条 本指引自公布之日起施行。", block_id="b6"),
            ],
            {
                "document_title": "公开募集证券投资基金投资信用衍生品指引",
                "issuing_authority": "中国证券监督管理委员会",
            },
        )
        clean_front_matter(document)
        values = [block.text for block in document.blocks]
        self.assertNotIn("证监会公告〔2019〕第1号", values)
        self.assertFalse(any(value.startswith("现公布") for value in values))
        self.assertIn("第十一条 本指引自公布之日起施行。", values)
        self.assertEqual(document.metadata["document_number"], "证监会公告〔2019〕第1号")
        self.assertEqual(document.cleaning["status"], "changed")

    def test_notice_body_and_referenced_document_number_are_preserved(self):
        document = ParsedDocument(
            Path("关于加强场外衍生品业务自律管理的通知.docx"),
            "docx",
            [
                SourceBlock("中证协发〔2017〕123号", block_id="b1"),
                SourceBlock("各证券公司：", block_id="b2"),
                SourceBlock("现将有关事项通知如下：", block_id="b3"),
                SourceBlock("一、执行中证协发〔2016〕8号文件。", block_id="b4"),
            ],
            {
                "document_title": "关于加强场外衍生品业务自律管理的通知",
                "issuing_authority": "中国证券业协会",
            },
        )
        clean_front_matter(document)
        values = [block.text for block in document.blocks]
        self.assertEqual(values[0], "各证券公司：")
        self.assertIn("现将有关事项通知如下：", values)
        self.assertIn("一、执行中证协发〔2016〕8号文件。", values)

    def test_split_guide_cover_is_metadata_not_chunk_body(self):
        document = ParsedDocument(
            Path("上交所期货公司股票期权经纪业务指南（2026年修订）.docx"),
            "docx",
            [
                SourceBlock("附件2", block_id="b1"),
                SourceBlock("期货公司股票期权", block_id="b2"),
                SourceBlock("经纪业务指南", block_id="b3"),
                SourceBlock("上海证券交易所", block_id="b4"),
                SourceBlock("2026年6月", block_id="b5"),
                SourceBlock("说明及声明", block_id="b6"),
                SourceBlock("本指南供期货公司开展股票期权经纪业务时参考。", block_id="b7"),
                SourceBlock("第一章 总体要求", block_id="b8"),
            ],
            {
                "document_title": "上海证券交易所期货公司股票期权经纪业务指南（2026年修订）",
                "issuing_authority": "上海证券交易所",
            },
        )
        clean_front_matter(document)
        values = [block.text for block in document.blocks]
        self.assertEqual(
            values,
            [
                "说明及声明",
                "本指南供期货公司开展股票期权经纪业务时参考。",
                "第一章 总体要求",
            ],
        )
        self.assertEqual(
            set(document.cleaning["rule_hits"]),
            {"cover_attachment_label", "cover_authority", "cover_document_title", "cover_month"},
        )

    def test_multiple_guide_cover_titles_are_removed_in_one_pass_and_cleaning_is_idempotent(self):
        document = ParsedDocument(
            Path("上交所证券公司股票期权经纪业务指南（2026年修订）.docx"),
            "docx",
            [
                SourceBlock("证券公司股票期权", block_id="b1"),
                SourceBlock("经纪业务指南", block_id="b2"),
                SourceBlock("上海证券交易所", block_id="b3"),
                SourceBlock("2026年6月", block_id="b4"),
                SourceBlock("证券公司股票期权经纪业务指南", block_id="b5"),
                SourceBlock("说明及声明", block_id="b6"),
                SourceBlock("第一章 总则", block_id="b7"),
                SourceBlock("第一条 本指南适用于证券公司。", block_id="b8"),
            ],
            {
                "document_title": "上海证券交易所证券公司股票期权经纪业务指南（2026年修订）",
                "issuing_authority": "上海证券交易所",
            },
        )

        clean_front_matter(document)
        once = [block.text for block in document.blocks]
        chars_after = document.cleaning["chars_after"]
        clean_front_matter(document)

        self.assertEqual(once, ["说明及声明", "第一章 总则", "第一条 本指南适用于证券公司。"])
        self.assertEqual([block.text for block in document.blocks], once)
        self.assertEqual(document.cleaning["chars_after"], chars_after)

    def test_split_publication_page_is_removed_but_formal_effective_article_is_preserved(self):
        document = ParsedDocument(
            Path("远期利率协议业务管理规定.doc"),
            "doc",
            [
                SourceBlock("中国人民银行公告", block_id="b1"),
                SourceBlock("〔2007〕第 20号", block_id="b2"),
                SourceBlock(
                    "中国人民银行制定了《远期利率协议业务管理规定》，现予公布。",
                    block_id="b3",
                ),
                SourceBlock("中国人民银行", block_id="b4"),
                SourceBlock("二〇〇七年九月二十九日", block_id="b5"),
                SourceBlock("远期利率协议业务管理规定", block_id="b6"),
                SourceBlock("第一条 为规范远期利率协议业务，制定本规定。", block_id="b7"),
                SourceBlock("第二十条 本规定自2007年11月1日起施行。", block_id="b8"),
            ],
            {
                "document_title": "中国人民银行公告[2007]第20号：远期利率协议业务管理规定",
                "issuing_authority": "中国人民银行",
            },
        )

        clean_front_matter(document)

        self.assertEqual(
            [block.text for block in document.blocks],
            ["第一条 为规范远期利率协议业务，制定本规定。", "第二十条 本规定自2007年11月1日起施行。"],
        )
        self.assertEqual(document.metadata["document_number"], "中国人民银行公告〔2007〕第20号")

    def test_substantive_notice_preamble_is_not_treated_as_publication_wrapper(self):
        document = ParsedDocument(
            Path("实施指引通知.docx"),
            "docx",
            [
                SourceBlock(
                    "我会组织起草了《实施指引》，现予发布，自即日起实施。现就有关事项通知如下：",
                    block_id="b1",
                ),
                SourceBlock("一、各机构应当遵照执行。", block_id="b2"),
            ],
            {"document_title": "关于发布实施指引的通知", "issuing_authority": "中国证券业协会"},
        )

        clean_front_matter(document)

        self.assertEqual(len(document.blocks), 2)

    def test_trailing_source_links_are_removed_after_final_article(self):
        document = ParsedDocument(
            Path("测试办法.doc"),
            "doc",
            [
                SourceBlock("第一条 正文。", block_id="b1"),
                SourceBlock("第二条 本办法自公布之日起施行。", block_id="b2"),
                SourceBlock("有关部门负责人就《测试办法》答记者问", block_id="b3"),
                SourceBlock("https://example.gov.cn/interview", block_id="b4"),
                SourceBlock("有关部门发布《测试办法》", block_id="b5"),
                SourceBlock("https://example.gov.cn/release", block_id="b6"),
            ],
            {"document_title": "测试办法", "issuing_authority": "测试机关"},
        )
        clean_front_matter(document)
        self.assertEqual(
            [block.text for block in document.blocks],
            ["第一条 正文。", "第二条 本办法自公布之日起施行。"],
        )
        self.assertIn("trailing_source_reference", document.cleaning["rule_hits"])

    def test_doc_extension_with_ooxml_signature_uses_docx_parser(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "扩展名错误的指南.doc"
            source = Path(directory) / "source.docx"
            word = Document()
            word.add_paragraph("测试业务指南")
            word.add_paragraph("第一条 正文。")
            word.save(source)
            path.write_bytes(source.read_bytes())
            parsed = parse_file(path)
        self.assertEqual(parsed.source_type, "docx")
        self.assertTrue(any("第一条 正文" in block.text for block in parsed.blocks))

    def test_front_page_timestamp_is_removed_but_article_date_is_kept(self):
        blocks = [
            SourceBlock("测试定义文件", block_id="b1"),
            SourceBlock("时间:2014-08-22", block_id="b2"),
            SourceBlock("声明", block_id="b3"),
            SourceBlock("第一条 自2026-11-16起施行。", block_id="b4"),
        ]
        filtered, removed = strip_front_page_metadata(blocks)
        self.assertEqual(removed, 1)
        self.assertEqual([block.block_id for block in filtered], ["b1", "b3", "b4"])

    def test_quoted_legal_basis_is_not_document_title(self):
        blocks = [SourceBlock("第一条 根据《中华人民共和国证券法》制定本办法。", block_id="b1")]
        metadata = infer_metadata(blocks, Path("证券公司操作风险管理指引.docx"))
        self.assertEqual(metadata["document_title"], "证券公司操作风险管理指引")
        self.assertEqual(metadata["document_title_source"], "filename")

    def test_source_prefix_is_removed_from_filename_title(self):
        metadata = infer_metadata([], Path("19_GFEX_广州期货交易所_广州期货交易所碳酸锂期货、期权业务细则.docx"))
        self.assertEqual(metadata["document_title"], "广州期货交易所碳酸锂期货、期权业务细则")

    def test_toc_entries_are_removed_as_a_group(self):
        lines = ["目 录", "第一章 总则 1", "第二章 证券发行 2", "第三章 证券交易 7", "中华人民共和国证券法"]
        filtered, removed = remove_toc_entries(lines)
        self.assertEqual(removed, 4)
        self.assertEqual(filtered, ["中华人民共和国证券法"])

    def test_non_chapter_toc_entries_are_removed(self):
        lines = ["目录", "一、总则 ................................................................ 2", "二、估值处理 ........ 3", "三、会计处理 ........ 4", "正文"]
        filtered, removed = remove_toc_entries(lines)
        self.assertEqual(removed, 4)
        self.assertEqual(filtered, ["正文"])

    def test_cross_page_sentence_and_word_are_joined(self):
        blocks = [
            SourceBlock("第五条 并至少每", page=1, block_id="b1"),
            SourceBlock("季度将管理情况书面报告。", page=2, block_id="b2"),
            SourceBlock("第六条 新条文。", page=2, block_id="b3"),
        ]
        merged, count = merge_cross_page_paragraphs(blocks)
        self.assertEqual(count, 1)
        self.assertEqual(merged[0].text, "第五条 并至少每季度将管理情况书面报告。")
        self.assertEqual(merged[1].text, "第六条 新条文。")

    def test_cross_page_numeric_continuation_is_not_mistaken_for_heading(self):
        blocks = [
            SourceBlock("家庭金融资产不低于", page=1, block_id="b1"),
            SourceBlock("500 万元,或者近 3 年本人年均收入不低于 40 万元;", page=2, block_id="b2"),
        ]
        merged, count = merge_cross_page_paragraphs(blocks)
        self.assertEqual(count, 1)
        self.assertIn("不低于500 万元", merged[0].text)

    def test_cross_page_join_stops_at_new_legal_structure(self):
        blocks = [
            SourceBlock("前一页末尾无句号", page=1, block_id="b1"),
            SourceBlock("（二）新的分项。", page=2, block_id="b2"),
        ]
        merged, count = merge_cross_page_paragraphs(blocks)
        self.assertEqual(count, 0)
        self.assertEqual(len(merged), 2)

    def test_cross_page_join_stops_at_top_level_guide_heading(self):
        blocks = [
            SourceBlock("证券投资基金投资信用衍生品估值指引(试行)", page=1, block_id="b1"),
            SourceBlock("一、总则", page=2, block_id="b2"),
            SourceBlock("(一)为规范基金投资信用衍生品估值行为。", page=2, block_id="b3"),
        ]
        merged, count = merge_cross_page_paragraphs(blocks)
        self.assertEqual(count, 0)
        self.assertEqual([block.text for block in merged], [block.text for block in blocks])

    def test_pdf_semantic_table_requires_two_nonempty_columns(self):
        note_box = [["", "注：这是一行说明", ""], ["", "继续说明", ""]]
        self.assertEqual(normalize_semantic_table(note_box), [])
        table = [["项目", "", "比例"], ["净资本", "", "100%"]]
        self.assertEqual(normalize_semantic_table(table), [["项目", "", "比例"], ["净资本", "", "100%"]])

    def test_pdf_table_cell_visual_wraps_are_collapsed(self):
        self.assertEqual(normalize_pdf_table_cell("债\n券、票据"), "债券、票据")
        self.assertEqual(normalize_pdf_table_cell("Net-to-\nGross Ratio"), "Net-to-Gross Ratio")
        self.assertEqual(normalize_pdf_table_cell("未来 30\n日现金流出"), "未来 30日现金流出")

    def test_semantic_table_separates_note_rows_and_preserves_order(self):
        rows = [
            ["注：这是表格前的较长说明文字，不应被当作表格单元格。", "", ""],
            ["说明的续行内容。", "", ""],
            ["受让方风险敞口", "=", "市场价值"],
            ["数值", "+", "独立金额"],
            ["由此可见，后续的较长解释文字也不是表格的数据行。", "", ""],
        ]
        segments = split_semantic_table_content(rows)
        self.assertEqual([kind for kind, _ in segments], ["text", "table", "text"])
        self.assertIn("说明的续行内容", segments[0][1])
        self.assertEqual(len(segments[1][1]), 2)
        self.assertIn("后续的较长解释", segments[2][1])

    def test_word_toc_is_removed_before_repeated_first_chapter(self):
        blocks = [
            SourceBlock("目录", block_id="b1"),
            SourceBlock("第一章 总则", block_id="b2"),
            SourceBlock("第二章 业务", block_id="b3"),
            SourceBlock("第一章 总则", block_id="b4"),
            SourceBlock("第一条 正文。", block_id="b5"),
        ]
        filtered, removed = strip_word_toc(blocks)
        self.assertEqual(removed, 3)
        self.assertEqual([block.block_id for block in filtered], ["b4", "b5"])

    def test_styled_word_toc_does_not_delete_chapter_based_guide_body(self):
        blocks = [
            SourceBlock("附件", style="Normal", block_id="b1"),
            SourceBlock("目录", style="Normal", block_id="b2"),
            SourceBlock("第一章 总体要求 1", style="toc 1", block_id="b3"),
            SourceBlock("一、概述 1", style="toc 2", block_id="b4"),
            SourceBlock("说明及声明", style="List Paragraph", block_id="b5"),
            SourceBlock("为便于做市商开展业务，制定本指南。", style="Normal", block_id="b6"),
            SourceBlock("第一章 总体要求", style="List Paragraph", block_id="b7"),
            SourceBlock("一、概述", style="Heading 2", block_id="b8"),
            SourceBlock("本所采用竞争性做市商制度。", style="Normal", block_id="b9"),
            SourceBlock("附件4：做市协议", style="Normal", block_id="b10"),
            SourceBlock("第二条 协议正文。", style="Normal", block_id="b11"),
        ]
        filtered, removed = strip_word_toc(blocks)
        self.assertEqual(removed, 3)
        self.assertEqual([block.block_id for block in filtered], ["b1", "b5", "b6", "b7", "b8", "b9", "b10", "b11"])

    def test_unstyled_guide_toc_is_removed_before_repeated_declaration(self):
        blocks = [
            SourceBlock("目 录", block_id="b1"),
            SourceBlock("说明及声明", block_id="b2"),
            SourceBlock("第一章 总体要求 1", block_id="b3"),
            SourceBlock("一、组织架构 2", block_id="b4"),
            SourceBlock("第二章 业务申请 6", block_id="b5"),
            SourceBlock("说明及声明", block_id="b6"),
            SourceBlock("本指南供业务办理时参考。", block_id="b7"),
            SourceBlock("第一章 总体要求", block_id="b8"),
        ]
        filtered, removed = strip_word_toc(blocks)
        self.assertEqual(removed, 5)
        self.assertEqual([block.block_id for block in filtered], ["b6", "b7", "b8"])

    def test_spaced_page_number_is_detected(self):
        self.assertTrue(is_page_number("— 1 0 —"))
        self.assertTrue(is_page_number("III"))

    def test_symbol_private_use_characters_are_normalized(self):
        self.assertEqual(clean_text("A\uf03dB\uf02bC\uf0b4D"), "A=B+C×D")
        self.assertNotIn("\uf03d", clean_text("A\uf03dB"))

    def test_wrapped_decimal_is_repaired_without_joining_list_items(self):
        self.assertEqual(clean_text("1.\n1 权益类衍生品交易"), "1.1 权益类衍生品交易")
        self.assertEqual(clean_text("5.\n625%—6.25%"), "5.625%—6.25%")
        self.assertEqual(clean_text("7.\n2.1交易条款"), "7.2.1交易条款")
        self.assertEqual(clean_text("1.\n2. 第二项"), "1.\n2. 第二项")

    def test_docx_smart_tag_text_is_not_dropped(self):
        document = Document()
        paragraph = document.add_paragraph("Shibor O/N:")
        smart_tag = OxmlElement("w:smartTag")
        run = OxmlElement("w:r")
        text = OxmlElement("w:t")
        text.text = "1M、3M、6M、9M、1Y"
        run.append(text)
        smart_tag.append(run)
        paragraph._p.append(smart_tag)
        wrapped = Paragraph(paragraph._p, paragraph._parent)
        self.assertEqual(paragraph_xml_text(wrapped), "Shibor O/N:1M、3M、6M、9M、1Y")

    def test_legacy_doc_toc_cleanup_keeps_body_articles(self):
        text = """测试定义文件\n目录\n第一条 通用定义 1\n第二条 日期定义 2\n第一条 通用定义\n正文\n第二条 日期定义\n正文二"""
        blocks, warnings = _blocks_from_plain_text(text)
        values = [block.text for block in blocks]
        self.assertEqual(values, ["测试定义文件", "第一条 通用定义", "正文", "第二条 日期定义", "正文二"])
        self.assertTrue(any("目录" in warning for warning in warnings))

    def test_legacy_doc_toc_cleanup_preserves_preface_before_first_article(self):
        text = """主协议\n目录\n第一条 协议构成 1\n第二条 支付义务 2\n为明确交易双方的权利和义务，双方签署本主协议。\n第一条 协议构成\n正文"""
        blocks, _ = _blocks_from_plain_text(text)
        self.assertEqual(
            [block.text for block in blocks],
            ["主协议", "为明确交易双方的权利和义务,双方签署本主协议。", "第一条 协议构成", "正文"],
        )

    def test_rate_definition_fraction_is_restored_from_verified_pdf_layout(self):
        blocks = [SourceBlock(
            "贴现因子适用如下公式:λ =1 + ( r ×N / D )其中,λ 是贴现因子。",
            page=16,
            block_id="b1",
        )]
        changed = apply_verified_formula_overrides(Path("中国银行间市场利率衍生产品交易定义文件（2022年版）.pdf"), blocks)
        self.assertEqual(changed, 1)
        self.assertIn("λ = 1 / [1 + (r × N / D)]", blocks[0].text)
        self.assertIn(r"\frac{1}{1+\frac{rN}{D}}", blocks[0].formula_data["latex"])

    def test_credit_valuation_formula_preserves_fraction_and_exponent_order(self):
        blocks = [SourceBlock(
            "(一)基于违约率的方法估值公式如下:*ln(1-D)乱码其中:V为估值。",
            page=6,
            block_id="b1",
        )]
        changed = apply_verified_formula_overrides(Path("证券投资基金投资信用衍生品估值指引(试行).pdf"), blocks)
        self.assertEqual(changed, 1)
        self.assertIn("exp((d / TY) × ln(1 - D))", blocks[0].text)
        self.assertIn("/ [1 + r_d × (d / TY)]", blocks[0].text)
        self.assertTrue(blocks[0].formula_data["latex_expressions"])
        self.assertIn(r"\frac", blocks[0].formula_data["latex_expressions"][0])

    def test_credit_default_probability_preserves_complement_event(self):
        blocks = [SourceBlock(
            "考虑 CRMW 风险下年化违约概率 D = P AB,且 P AB = P A| B ×P B。",
            page=10,
            block_id="b1",
        )]
        changed = apply_verified_formula_overrides(Path("证券投资基金投资信用衍生品估值指引(试行).pdf"), blocks)
        self.assertEqual(changed, 1)
        self.assertIn("P(A∩B̄) = P(A|B̄) × P(B̄)", blocks[0].text)
        self.assertEqual(len(blocks[0].formula_data["latex_expressions"]), 2)

    def test_credit_complex_formula_keeps_condition_separate_and_latex(self):
        blocks = [SourceBlock(
            "(四)估值公式如下:∫_{t}^{T_n}旧内容其中:T_j<t<T_{j+1}",
            page=8,
            block_id="b1",
        )]
        changed = apply_verified_formula_overrides(Path("证券投资基金投资信用衍生品估值指引(试行).pdf"), blocks)
        self.assertEqual(changed, 1)
        self.assertIn("注：T_j < t < T_{j+1}", blocks[0].text)
        self.assertIn(r"\int", blocks[0].formula_data["latex"])

    def test_legacy_doc_missing_ole_formula_is_marked_without_inventing_content(self):
        document = ParsedDocument(
            Path("商业银行资本管理办法.doc"),
            "doc",
            [
                SourceBlock("第十九条 商业银行资本充足率计算公式为:", block_id="b1"),
                SourceBlock("第二十条 商业银行杠杆率计算公式为:", block_id="b2"),
                SourceBlock("第二十一条 商业银行总资本包括一级资本。", block_id="b3"),
                SourceBlock("第一百二十条 内部损失乘数计算公式为:", block_id="b4"),
                SourceBlock("其中:", block_id="b5"),
            ],
            {"document_title": "商业银行资本管理办法"},
        )
        rows = structured_blocks(document)
        missing = [row for row in rows if row.get("formula_data", {}).get("conversion_status") == "source_formula_not_extractable"]
        self.assertEqual([row["block_id"] for row in missing], ["b1", "b2", "b4"])
        self.assertTrue(all(not row["formula_data"]["expressions"] for row in missing))
        self.assertTrue(all("请对照原件" in row["parsing_warnings"][0] for row in missing))

    def test_legacy_doc_verified_ole_formulas_are_restored_from_original_layout(self):
        blocks = [
            SourceBlock("第十九条 商业银行资本充足率计算公式为:", block_id="b19"),
            SourceBlock("第二十条 商业银行杠杆率计算公式为:", block_id="b20"),
            SourceBlock("第一百二十条 内部损失乘数(ILM)是基于商业银行操作风险平均历史损失数据与业务指标部分的调整因子,计算公式为:", block_id="b120"),
        ]
        changed = apply_verified_doc_formula_overrides(Path("商业银行资本管理办法.doc"), blocks)
        self.assertEqual(changed, 3)
        self.assertEqual(len(blocks[0].formula_data["expressions"]), 3)
        self.assertIn("核心一级资本充足率", blocks[0].text)
        self.assertIn("调整后表内外资产余额", blocks[1].text)
        self.assertIn("(LC / BIC)^{0.8}", blocks[2].text)
        self.assertEqual(blocks[2].formula_data["conversion_status"], "verified_from_original_doc_layout")

    def test_official_html_footnote_marker_is_structured(self):
        parser = _OfficialBodyHTML()
        parser.feed('<div class="TRS_Editor"><p>第十七条① 正文。</p><p>①原第十七条已删除。</p></div>')
        self.assertEqual(parser.parts, ["第十七条 正文。", "修订注1：原第十七条已删除。"])


if __name__ == "__main__":
    unittest.main()
